"""
Document generator — builds a professional Microsoft Word (.docx) file
from the assembled document content using python-docx.
"""

import logging
import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import settings
from utils import generate_filename

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _set_heading_color(paragraph, r: int, g: int, b: int) -> None:
    """Apply an RGB colour to every run in a paragraph."""
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _add_horizontal_rule(document: Document) -> None:
    """Insert a thin horizontal line (bottom border on an empty paragraph)."""
    p = document.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _parse_sections(content: str) -> dict[str, str]:
    """
    Parse the flat document content string into named sections.

    Looks for common section markers produced by the execution / regeneration
    prompts. Falls back to treating the whole content as 'Main Content'.

    Args:
        content: Raw multi-line document content string.

    Returns:
        Ordered dict mapping section title → section text.
    """
    section_patterns = [
        ("TITLE", r"TITLE\s*:\s*(.+?)(?=\n[A-Z ]+:|$)"),
        ("EXECUTIVE SUMMARY", r"EXECUTIVE SUMMARY\s*:\s*(.*?)(?=\n[A-Z ]+:|$)"),
        ("OBJECTIVES", r"OBJECTIVES\s*:\s*(.*?)(?=\n[A-Z ]+:|$)"),
        ("MAIN CONTENT", r"MAIN CONTENT\s*:\s*(.*?)(?=\n[A-Z ]+:|$)"),
        ("RECOMMENDATIONS", r"RECOMMENDATIONS\s*:\s*(.*?)(?=\n[A-Z ]+:|$)"),
        ("CONCLUSION", r"CONCLUSION\s*:\s*(.*?)(?=\n[A-Z ]+:|$)"),
    ]

    parsed: dict[str, str] = {}
    for label, pattern in section_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            parsed[label] = match.group(1).strip()

    if not parsed:
        # No markers found — treat full content as main body
        parsed["MAIN CONTENT"] = content.strip()

    return parsed


def _add_bullet_lines(document: Document, text: str) -> None:
    """
    Add text lines as bullet points if the text contains list-like items,
    otherwise add as normal paragraph(s).

    Detects lines starting with -, *, •, or numbered (1. 2. etc.).
    """
    lines = text.split("\n")
    in_list = False
    para_buffer: list[str] = []

    def flush_buffer() -> None:
        nonlocal para_buffer
        if para_buffer:
            document.add_paragraph(" ".join(para_buffer))
            para_buffer = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_buffer()
            continue

        bullet_match = re.match(r"^[-*•]\s+(.+)$", stripped)
        numbered_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)

        if bullet_match or numbered_match:
            flush_buffer()
            item_text = (bullet_match or numbered_match).group(1)
            p = document.add_paragraph(item_text, style="List Bullet")
            in_list = True
        else:
            if in_list:
                in_list = False
            para_buffer.append(stripped)

    flush_buffer()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_document(
    document_type: str,
    request: str,
    section_outputs: dict[str, str],
    final_content: str,
) -> str:
    """
    Build and save a professional Word document.

    Strategy:
    1. Try to use the structured `final_content` (from reflection step) if it
       contains recognisable section markers.
    2. Fall back to `section_outputs` (raw executor outputs keyed by task name)
       and map them to canonical document sections.

    Args:
        document_type:   Title / type of the document.
        request:         Original user request (shown as a subtitle).
        section_outputs: Dict of task → output from the executor.
        final_content:   Assembled / possibly regenerated full content.

    Returns:
        Absolute file path of the saved .docx file.
    """
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- Title ---
    title_para = doc.add_heading(document_type, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title_para, 31, 73, 125)  # dark blue

    # --- Subtitle (request excerpt) ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(request[:200] + ("..." if len(request) > 200 else ""))
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(89, 89, 89)

    # --- Generation timestamp ---
    ts_para = doc.add_paragraph()
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_run = ts_para.add_run(f"Prepared: {datetime.utcnow().strftime('%B %d, %Y  %H:%M UTC')}")
    ts_run.font.size = Pt(9)
    ts_run.font.color.rgb = RGBColor(128, 128, 128)

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # --- Parse sections ---
    sections = _parse_sections(final_content)

    # If structured parsing produced content, use it; otherwise fall back to executor outputs
    # A stray label such as "Conclusion:" inside an executor output is not a
    # structured full document. Require at least two parsed sections.
    if len(sections) <= 1:
        sections = _build_sections_from_executor(section_outputs)

    canonical_order = [
        "EXECUTIVE SUMMARY",
        "OBJECTIVES",
        "MAIN CONTENT",
        "RECOMMENDATIONS",
        "CONCLUSION",
    ]

    for section_title in canonical_order:
        content_text = sections.get(section_title, "")
        if not content_text:
            continue

        heading = doc.add_heading(section_title.title(), level=1)
        _set_heading_color(heading, 31, 73, 125)
        _add_bullet_lines(doc, content_text)
        doc.add_paragraph()  # spacing

    # --- Footer ---
    _add_horizontal_rule(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Confidential")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(150, 150, 150)

    # --- Save ---
    filename = generate_filename(document_type)
    filepath = os.path.join(settings.OUTPUT_DIR, filename)
    os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
    doc.save(filepath)

    logger.info("Document saved: %s", filepath)
    return filepath


def _build_sections_from_executor(section_outputs: dict[str, str]) -> dict[str, str]:
    """
    Map executor task outputs to canonical document sections when the
    final_content string has no structured markers.

    Heuristic keyword mapping:
    - Tasks mentioning 'summary' / 'overview' / 'understand' → EXECUTIVE SUMMARY
    - Tasks mentioning 'objective' / 'goal' / 'scope'        → OBJECTIVES
    - Tasks mentioning 'recommend'                           → RECOMMENDATIONS
    - Tasks mentioning 'conclusion' / 'review'               → CONCLUSION
    - Everything else                                        → MAIN CONTENT
    """
    mapping: dict[str, list[str]] = {
        "EXECUTIVE SUMMARY": [],
        "OBJECTIVES": [],
        "MAIN CONTENT": [],
        "RECOMMENDATIONS": [],
        "CONCLUSION": [],
    }

    keyword_rules: list[tuple[str, list[str]]] = [
        ("OBJECTIVES", ["objective", "goal", "scope", "purpose"]),
        ("EXECUTIVE SUMMARY", ["summary", "overview", "understand", "analyze", "executive"]),
        ("RECOMMENDATIONS", ["recommend", "suggestion", "advise"]),
        ("CONCLUSION", ["conclusion", "review", "final", "closing"]),
    ]

    for task, output in section_outputs.items():
        task_lower = task.lower()
        conclusion_match = re.search(r"\bConclusion\s*:\s*(.+)$", output, re.IGNORECASE | re.DOTALL)
        if conclusion_match:
            mapping["CONCLUSION"].append(conclusion_match.group(1).strip())
            output = output[:conclusion_match.start()].rstrip()
        matched = False
        for section_key, keywords in keyword_rules:
            if any(kw in task_lower for kw in keywords):
                mapping[section_key].append(output)
                matched = True
                break
        if not matched:
            mapping["MAIN CONTENT"].append(output)

    if not mapping["CONCLUSION"]:
        mapping["CONCLUSION"].append(
            "Proceed through stakeholder validation and staged delivery, measure the agreed "
            "success criteria, and revisit assumptions at each decision checkpoint."
        )

    return {k: "\n\n".join(v) for k, v in mapping.items() if v}
